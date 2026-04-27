# Install library
!pip install python-docx

from docx import Document
from docx.shared import Pt

doc = Document()


def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text=""):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


add_heading('UJIAN TENGAH SEMESTER', 0)

add_paragraph('Topik: Analisis Big Data\n')
add_paragraph('Mata Kuliah IF25-40402 – Mahadata\n')

add_paragraph('Disusun oleh:')
add_paragraph('- Nama 1 – NIM')
add_paragraph('- Nama 2 – NIM')
add_paragraph('- Nama 3 – NIM\n')

add_paragraph('Dosen Pengampu: [Nama Dosen]')
add_paragraph('Program Studi Teknik Informatika')
add_paragraph('Institut Teknologi Sumatera')
add_paragraph('Semester Genap 2025/2026')

doc.add_page_break()

add_heading('BAGIAN 2 PENDAHULUAN', 1)

add_heading('2.1 Latar Belakang', 2)
add_paragraph('[Jelaskan pentingnya analitik big data: real-time analytics, ML at scale, dsb]')

add_heading('2.2 Tujuan Review', 2)
add_paragraph(
    '- Mengidentifikasi metode analitik big data\n'
    '- Membandingkan pendekatan antar jurnal\n'
    '- Mengevaluasi kelebihan dan kelemahan\n'
)

add_heading('2.3 Ruang Lingkup', 2)
add_paragraph('[Fokus pada stream processing, ML at scale, dll]')

add_heading('2.4 Metodologi Pencarian Jurnal', 2)
add_paragraph(
    'Database: IEEE, ACM, Springer, ScienceDirect\n'
    'Kata kunci: "big data" AND "stream processing" AND "real-time analytics"\n'
    'Kriteria: 2023–2026, peer-reviewed, relevan dengan analitik big data'
)

doc.add_page_break()

# =========================
# BAGIAN 3: REVIEW JURNAL
# =========================
add_heading('BAGIAN 3 REVIEW JURNAL', 1)

for i in range(1, 6):
    add_heading(f'3.{i} Jurnal {i}', 2)

    add_heading(f'3.{i}.1 Identitas Jurnal', 3)
    add_paragraph('Judul:\nPenulis:\nVenue:\nTahun:\nDOI:\nPublisher:')

    add_heading(f'3.{i}.2 Ringkasan Konten', 3)
    add_paragraph(
        'Problem Statement:\n\n'
        'Metode:\n\n'
        'Kontribusi:\n\n'
        'Dataset & Skala:\n'
        '- Ukuran data (GB/TB)\n'
        '- Jumlah node\n'
        '- Benchmark (TPC-DS, HiBench, dll)\n\n'
        'Hasil:\n'
        '- Accuracy\n'
        '- Latency\n'
        '- Throughput'
    )

    add_heading(f'3.{i}.3 Kekuatan', 3)
    add_paragraph(
        '- Skalabilitas terbukti secara eksperimen\n'
        '- Evaluasi pada dataset besar\n'
        '- Perbandingan dengan metode lain\n'
    )

    add_heading(f'3.{i}.4 Kelemahan', 3)
    add_paragraph(
        '- Asumsi tidak realistis\n'
        '- Skala tidak mencerminkan kondisi nyata\n'
        '- Tidak membahas cost/performance\n'
    )

    add_heading(f'3.{i}.5 Kaitan dengan Materi Kuliah', 3)
    add_paragraph(
        '- Distributed computing\n'
        '- Stream processing\n'
        '- Machine learning at scale'
    )

doc.add_page_break()

# =========================
# BAGIAN 4: ANALISIS KOMPARATIF
# =========================
add_heading('BAGIAN 4 ANALISIS KOMPARATIF', 1)

add_heading('4.1 Tabel Perbandingan', 2)

jumlah_jurnal = 5
table = doc.add_table(rows=1, cols=jumlah_jurnal + 1)
hdr_cells = table.rows[0].cells

headers = ['Aspek'] + [f'J{i}' for i in range(1, jumlah_jurnal + 1)]
for i, h in enumerate(headers):
    hdr_cells[i].text = h

aspek_list = [
    'Metode',
    'Hasil',
    'Arsitektur',
    'Skala Data',
    'Latency',
    'Throughput',
    'Scalability',
    'Kelebihan',
    'Kekurangan'
]

for aspek in aspek_list:
    row_cells = table.add_row().cells
    row_cells[0].text = aspek

add_heading('4.2 Perbandingan Pendekatan', 2)
add_paragraph('[Bandingkan batch vs stream vs hybrid]')

add_heading('4.3 Analisis Tren', 2)
add_paragraph('[Perkembangan real-time analytics, distributed ML, dll]')

add_heading('4.4 Jurnal Paling Berkontribusi', 2)
add_paragraph('[Pilih dan jelaskan dengan argumen kuat]')

add_heading('4.5 Research Gap', 2)
add_paragraph('[Celah penelitian yang belum terselesaikan]')

doc.add_page_break()

# =========================
# BAGIAN 5: KESIMPULAN
# =========================
add_heading('BAGIAN 5 KESIMPULAN DAN SARAN', 1)

add_heading('5.1 Kesimpulan', 2)
add_paragraph('[Rangkuman temuan utama]')

add_heading('5.2 Saran', 2)
add_paragraph('[Arah penelitian selanjutnya + implikasi industri]')

doc.add_page_break()

# =========================
# BAGIAN 6: DAFTAR PUSTAKA
# =========================
add_heading('BAGIAN 6 DAFTAR PUSTAKA', 1)
add_paragraph(
    '[1] Author, "Title," Journal, vol., no., year, DOI.\n'
    '[2] Author, "Title," Conference, year, DOI.\n'
)

doc.add_page_break()

# =========================
# BAGIAN 7: LAMPIRAN
# =========================
add_heading('BAGIAN 7 LAMPIRAN', 1)

add_heading('7.1 Pembagian Tugas', 2)
table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nama'
hdr_cells[1].text = 'Tugas'
hdr_cells[2].text = 'Kontribusi (%)'

add_heading('7.2 Screenshot Jurnal', 2)
add_paragraph('[Masukkan halaman pertama jurnal]')

# Save
file_path = '/content/UTS_Mahadata_Analisis_Big_Data.docx'
doc.save(file_path)

file_path
